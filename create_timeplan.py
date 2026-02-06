from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Create document
doc = Document()

# Add title
title = doc.add_heading('Time Plan', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Define table data
data = [
    ["Task", "Duration (Days)", "Start", "Finish", "Resource Name"],
    ["Graduation Project", "246", "05/10/2025", "11/06/2026", ""],
    ["Research & Topic Selection", "8", "05/10/2025", "13/10/2025", "All Team Members"],
    ["Business Requirements", "23", "14/10/2025", "08/11/2025", "Nada, Sandy"],
    ["Tools Installation", "3", "14/10/2025", "17/10/2025", "Yahia, Ahmed"],
    ["Technology Study", "21", "18/10/2025", "10/11/2025", "All Team Members"],
    ["System Analysis", "21", "", "", ""],
    ["Use Case", "21", "17/11/2025", "08/12/2025", "Nada"],
    ["Use Case Scenarios", "21", "17/11/2025", "08/12/2025", "Sandy"],
    ["DFD", "21", "17/11/2025", "08/12/2025", "Ali"],
    ["Design", "45", "", "", ""],
    ["UI (Figma)", "30", "17/11/2025", "17/12/2025", "Noor, Yahia"],
    ["Sequence Diagrams", "30", "17/11/2025", "15/12/2025", "Noor, Yahia"],
    ["ERD Design", "14", "01/12/2025", "15/12/2025", "Ahmed, Ziad"],
    ["Database Schema", "10", "16/12/2025", "26/12/2025", "Ahmed, Ziad"],
    ["Class Diagrams", "30", "17/11/2025", "15/12/2025", "Ahmed, Ziad"],
    ["System Architecture", "12", "09/05/2026", "21/05/2026", "Ahmed"],
    ["Programming", "163", "", "", ""],
    ["Frontend Implementation", "42", "18/12/2025", "01/02/2026", "Yahia, Nada"],
    ["Backend Implementation", "90", "09/02/2026", "09/05/2026", "Ahmed, Ziad"],
    ["Database Implementation", "14", "09/02/2026", "23/02/2026", "Ahmed, Ziad"],
    ["API Development", "70", "24/02/2026", "05/05/2026", "Ahmed, Ziad"],
    ["AI Implementation", "121", "18/12/2025", "19/04/2026", "Ali, Mohamed, Sandy"],
    ["AI API Integration", "20", "20/04/2026", "09/05/2026", "Ali, Ahmed"],
    ["Testing", "10", "", "", ""],
    ["Frontend Testing", "10", "09/05/2026", "19/05/2026", "Yahia, Nada"],
    ["Backend Testing", "10", "09/05/2026", "19/05/2026", "Ahmed, Ziad"],
    ["AI Testing", "10", "09/05/2026", "19/05/2026", "Ali, Mohamed, Sandy"],
    ["Documentation", "198", "27/12/2025", "08/06/2026", "Sandy"],
    ["Meetings (Review)", "3", "09/06/2026", "11/06/2026", "All Team Members"],
]

# Bold/header rows (phase headers)
bold_rows = [0, 1, 6, 10, 17, 24]

# Create table
table = doc.add_table(rows=len(data), cols=5)
table.style = 'Table Grid'

# Set column widths
widths = [Cm(5), Cm(3), Cm(2.5), Cm(2.5), Cm(5)]

# Populate table
for i, row_data in enumerate(data):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        
        # Set cell width
        cell.width = widths[j]
        
        # Format paragraph
        para = cell.paragraphs[0]
        
        # Center align duration, start, finish columns
        if j in [1, 2, 3]:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Bold for header and phase rows
        if i in bold_rows:
            for run in para.runs:
                run.bold = True
        
        # Set font size
        for run in para.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'

# Make header row have gray background
for cell in table.rows[0].cells:
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
    cell._tc.get_or_add_tcPr().append(shading)

# Save document
doc.save('Docs/TimePlan_Updated.docx')
print("Word document created successfully: Docs/TimePlan_Updated.docx")
